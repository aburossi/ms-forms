import streamlit as st
from openai import OpenAI
import json
import random
import PyPDF2
import docx
from PIL import Image
import io
import base64
import logging
import re
from pdf2image import convert_from_bytes  # Added for PDF to image conversion
import streamlit.components.v1 as components  # Added for embedding HTML components

# Set up logging for better error tracking
logging.basicConfig(level=logging.INFO)

# Initialize OpenAI client (moved outside main to maintain global scope)
def initialize_openai(api_key):
    return OpenAI(api_key=api_key)

# List of supported languages
LANGUAGES = {
    "Deutsch": "German",
    "Englisch": "English",
    "Französisch": "French",
    "Italienisch": "Italian",
    "Spanisch": "Spanish"
}

# List of supported file types
SUPPORTED_FILE_TYPES = ["pdf", "docx", "jpg", "jpeg", "png"]

# Caching functions for performance
@st.cache_data
def extract_text_from_pdf(file):
    """Extract text from PDF using PyPDF2."""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        return ""

@st.cache_data
def extract_text_from_docx(file):
    """Extract text from DOCX file."""
    try:
        doc = docx.Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        return ""

@st.cache_data
def convert_pdf_to_images(file):
    """Convert PDF pages to images."""
    try:
        images = convert_from_bytes(file.read())
        return images
    except Exception as e:
        logging.error(f"Error converting PDF to images: {e}")
        return []

def process_image(_image):
    """Process and resize an image to reduce memory footprint."""
    try:
        if isinstance(_image, (str, bytes)):
            img = Image.open(io.BytesIO(base64.b64decode(_image) if isinstance(_image, str) else _image))
        elif isinstance(_image, Image.Image):
            img = _image
        else:
            img = Image.open(_image)

        # Convert to RGB mode if it's not
        if img.mode != 'RGB':
            img = img.convert('RGB')

        # Resize if the image is too large
        max_size = 1000  # Reduced max size to reduce memory consumption
        if max(img.size) > max_size:
            img.thumbnail((max_size, max_size))

        # Save to bytes
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='JPEG')
        img_byte_arr = img_byte_arr.getvalue()

        return base64.b64encode(img_byte_arr).decode('utf-8')
    except Exception as e:
        logging.error(f"Error processing image: {e}")
        return None

def get_chatgpt_response(prompt, num_questions, language, image=None):
    """Get response from OpenAI API with error handling."""
    try:
        system_prompt = f"""
        You are an expert in education and quiz creation. Generate {num_questions} quiz questions based on the provided content.
        The questions and answers should be in {language}.
        
        Return your response in the following JSON format EXACTLY:
        {{
            "questions": [
                {{
                    "question": "Your question text here",
                    "correct_answer": "The correct answer",
                    "incorrect_answers": ["Wrong answer 1", "Wrong answer 2"]
                }},
                ... (more questions)
            ]
        }}

        Mix Multiple Choice and Fill in the Blank questions:
        - For Multiple Choice: Write a complete question
        - For Fill in the Blank: Include _____ in the question text
        
        Make all answers and questions clear and unambiguous.
        IMPORTANT: 
        - Ensure your response is valid JSON that can be parsed
        - ALL text must be in {language}
        - Make sure the questions are culturally appropriate for {language}-speaking regions
        """

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ]

        if image:
            base64_image = process_image(image)
            if base64_image:
                # Embed the image in the prompt using a data URL
                image_data_url = f"data:image/jpeg;base64,{base64_image}"
                messages[1]["content"] = f"{prompt}\n![Image]({image_data_url})"
            else:
                st.warning("Bildverarbeitung fehlgeschlagen. Fortfahren ohne das Bild.")

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=8000,
            temperature=0.5
        )

        # Get the response content
        response_text = response.choices[0].message.content.strip()
        
        # Try to find JSON content within the response
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            json_str = json_match.group(0)
            # Parse and validate JSON structure
            questions_data = json.loads(json_str)
            if "questions" not in questions_data:
                # Convert old format to new format if necessary
                if isinstance(questions_data, list):
                    questions_data = {"questions": questions_data}
            return questions_data["questions"]
        else:
            raise ValueError("Kein gültiges JSON in der Antwort gefunden")

    except json.JSONDecodeError as e:
        st.error(f"Fehler beim Parsen der JSON-Antwort: {str(e)}")
        logging.error(f"Rohantwort: {response_text}")
        return None
    except Exception as e:
        st.error(f"Fehler bei der Kommunikation mit der OpenAI API: {str(e)}")
        logging.error(f"Fehler: {e}")
        return None

def generate_word_document(questions):
    """Generate a Word document from the quiz questions."""
    try:
        doc = docx.Document()
        
        for i, q in enumerate(questions, start=1):
            # Validate question structure
            if not isinstance(q, dict) or not all(k in q for k in ['question', 'correct_answer', 'incorrect_answers']):
                continue
                
            # Replace 'ß' with 'ss' in question and answers
            question_text = q['question'].replace('ß', 'ss')
            correct_answer = q['correct_answer'].replace('ß', 'ss')
            incorrect_answers = [ans.replace('ß', 'ss') for ans in q['incorrect_answers']]
            
            # Combine and randomize answers
            all_answers = incorrect_answers + [correct_answer]
            random.shuffle(all_answers)
            
            # Determine correct answer label
            answer_labels = ['A', 'B', 'C']
            correct_label = next(label for label, ans in zip(answer_labels, all_answers) 
                               if ans == correct_answer)
            
            # Add question
            doc.add_paragraph(f"{i}. {question_text}")
            
            # Add options
            for label, option in zip(answer_labels, all_answers):
                doc.add_paragraph(f"{label}. {option}")
            
            # Add correct answer
            doc.add_paragraph(f"Antwort: {correct_label}")
            
            # Add spacing
            doc.add_paragraph()
        
        # Save document to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    except Exception as e:
        st.error(f"Fehler beim Erstellen des Word-Dokuments: {str(e)}")
        logging.error(f"Dokumenterstellungsfehler: {str(e)}")
        return None

def process_images(images, language, num_questions):
    """Process uploaded images and generate questions."""
    all_questions = []
    for idx, image in enumerate(images):
        st.image(image, caption=f'Seite {idx+1}', use_column_width=True)
        
        # Optional: Allow users to provide specific prompts per image
        user_input = st.text_area(f"Geben Sie Ihren Text ein oder beschreiben Sie den Inhalt für Seite {idx+1}:", key=f"text_area_{idx}")
        learning_goals = st.text_area(f"Lernziele für Seite {idx+1} (Optional):", key=f"learning_goals_{idx}")
        
        if st.button(f"Fragen für Seite {idx+1} generieren", key=f"generate_button_{idx}"):
            if user_input:
                prompt = f"{user_input}\nLernziele: {learning_goals}"
            else:
                prompt = f"Extrahierter Inhalt aus Bild {idx+1}.\nLernziele: {learning_goals}"
            
            with st.spinner(f"Generiere Fragen für Seite {idx+1}..."):
                questions = get_chatgpt_response(prompt, num_questions, language, image=image)
                if questions:
                    all_questions.extend(questions)
                    st.success(f"Fragen für Seite {idx+1} erfolgreich generiert!")
                else:
                    st.error(f"Fehler beim Generieren der Fragen für Seite {idx+1}.")
    return all_questions

def main():
    """Main function for the Streamlit app."""
    st.title("📝 Microsoft Forms Quiz Generator")

    # Sidebar Instructions
    with st.sidebar:
        st.header("❗ **So verwenden Sie diese App**")
        st.markdown("""
        1. **Geben Sie Ihren OpenAI-API-Schlüssel ein**: Erhalten Sie Ihren API-Schlüssel von [OpenAI](https://platform.openai.com/account/api-keys).
        2. **Datei hochladen oder Text eingeben**: Sie können eine PDF-, DOCX- oder Bilddatei hochladen oder direkt Text eingeben.
        3. **Lernziele angeben** (optional): Definieren Sie die Lernziele, die die Quizfragen abdecken sollen.
        4. **Sprache für Fragen wählen**: Wählen Sie die Sprache, in der die Quizfragen generiert werden sollen.
        5. **Anzahl der Fragen auswählen**: Bestimmen Sie, wie viele Fragen generiert werden sollen.
        6. **Quiz generieren**: Klicken Sie auf "Quiz generieren", um das Quiz zu erstellen und als Word-Dokument herunterzuladen.
        """)
        
        components.html("""
            <iframe width="100%" height="180" src="https://www.youtube.com/embed/OB99E7Y1cMA" 
            title="Demo-Video auf Deutsch" frameborder="0" allow="accelerometer; autoplay; 
            clipboard-write; encrypted-media; gyroscope; picture-in-picture" allowfullscreen>
            </iframe>
        """, height=180)
        
        st.markdown("---")
        st.header("📜 Lizenz")
        st.markdown("Diese Anwendung steht unter der [MIT-Lizenz](https://opensource.org/licenses/MIT).")
        
        st.header("💬 Kontakt")
        st.markdown("**Kontakt**: [Pietro](mailto:pietro.rossi@bbw.ch)")

    # Language selection with radio buttons
    st.subheader("Wählen Sie die Ausgabesprache für die generierten Fragen:")
    selected_language_key = st.radio("Wählen Sie die Ausgabesprache:", list(LANGUAGES.keys()), index=0)
    selected_language = LANGUAGES[selected_language_key]

    # File upload
    uploaded_file = st.file_uploader("Datei hochladen (PDF, DOCX oder Bild)", type=SUPPORTED_FILE_TYPES)

    content = ""
    image_content = None
    images = []

    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            content = extract_text_from_pdf(uploaded_file)
            if content:
                st.success("Text aus PDF extrahiert. Sie können ihn im folgenden Textbereich bearbeiten.")
            else:
                st.warning("Kein Text im PDF gefunden. Konvertiere PDF-Seiten in Bilder.")
                images = convert_pdf_to_images(uploaded_file)
                if images:
                    st.success("PDF in Bilder konvertiert. Sie können jetzt Fragen für jede Seite generieren.")
                else:
                    st.error("Fehler beim Konvertieren des PDFs in Bilder.")
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            content = extract_text_from_docx(uploaded_file)
            if content:
                st.success("Text aus DOCX extrahiert. Sie können ihn im folgenden Textbereich bearbeiten.")
            else:
                st.error("Fehler beim Extrahieren des Textes aus DOCX.")
        elif uploaded_file.type.startswith('image/'):
            image_content = Image.open(uploaded_file)
            st.image(image_content, caption='Hochgeladenes Bild', use_column_width=True)
            st.success("Bild erfolgreich hochgeladen.")
        else:
            st.error("Nicht unterstützter Dateityp. Bitte laden Sie eine PDF-, DOCX- oder Bilddatei hoch.")

    # Text input area
    user_input = st.text_area("Geben Sie Ihren Text ein oder beschreiben Sie das Bild:", 
                             value=content if content else "")
    
    # Learning goals input
    learning_goals = st.text_area("Lernziele (optional):")

    # Number of questions slider
    num_questions = st.slider("Anzahl der zu generierenden Fragen:", 
                            min_value=1, 
                            max_value=20, 
                            value=5)

    # API key input
    api_key = st.text_input("Geben Sie Ihren OpenAI API-Schlüssel ein:", type="password")

    # Generate button
    if st.button("Quiz generieren"):
        if not api_key:
            st.error("Bitte geben Sie Ihren OpenAI API-Schlüssel ein, um fortzufahren.")
            return

        if not user_input and not uploaded_file:
            st.warning("Bitte geben Sie Text ein oder laden Sie eine Datei hoch.")
            return

        # Initialize OpenAI client
        global client
        client = initialize_openai(api_key)

        all_questions = []

        if uploaded_file and uploaded_file.type == "application/pdf" and not content:
            # If PDF was uploaded and text extraction failed, process images
            if images:
                all_questions = process_images(images, selected_language, num_questions)
            else:
                st.error("Keine Bilder zum Verarbeiten aus dem PDF.")
        elif uploaded_file and uploaded_file.type.startswith('image/'):
            # If an image was uploaded directly
            with st.spinner("Generiere Fragen aus dem hochgeladenen Bild..."):
                questions = get_chatgpt_response(user_input, num_questions, selected_language, image=image_content)
                if questions:
                    all_questions.extend(questions)
                    st.success("Fragen erfolgreich generiert!")
                else:
                    st.error("Fehler beim Generieren der Fragen aus dem Bild.")
        else:
            # Process text input
            prompt = f"""
            Thema/Inhalt: {user_input}
            Lernziele: {learning_goals}
            Ausgabesprache: {selected_language}
            
            Bitte generiere {num_questions} Quizfragen basierend auf diesem Inhalt, wobei Multiple Choice und Lückentexte gemischt werden.
            Stelle sicher, dass deine Antwort im korrekten JSON-Format mit einem 'questions'-Array ist.
            Generiere alle Fragen und Antworten in {selected_language}.
            """
            with st.spinner("Generiere Quizfragen..."):
                questions = get_chatgpt_response(prompt, num_questions, selected_language)
                if questions:
                    all_questions.extend(questions)
                    st.success("Fragen erfolgreich generiert!")
                else:
                    st.error("Fehler beim Generieren der Fragen.")

        if all_questions:
            # Generate Word document
            doc_bytes = generate_word_document(all_questions)
            
            if doc_bytes:
                # Download button
                st.download_button(
                    label="Quiz-Dokument herunterladen",
                    data=doc_bytes,
                    file_name=f"Quiz_für_Microsoft_Forms_{selected_language}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Show preview of questions
                with st.expander("Vorschau der generierten Fragen"):
                    for i, q in enumerate(all_questions, 1):
                        st.markdown(f"**Frage {i}: {q['question']}**")
                        st.markdown(f"**Richtige Antwort:** {q['correct_answer']}")
                        st.markdown(f"**Falsche Antworten:** {', '.join(q['incorrect_answers'])}")
                
                # Show import instructions
                st.success("Quiz erfolgreich generiert! Laden Sie das Dokument herunter und importieren Sie es in Microsoft Forms.")
                st.markdown("Brauchen Sie Hilfe beim Importieren? Überprüfen Sie die Anleitungen und Tutorials in der Seitenleiste.")

if __name__ == "__main__":
    main()
